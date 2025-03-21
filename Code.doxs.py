from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def generer_sujet_docx(questions, numero_sujet, nom_fichier):
    doc = Document()

    # Paysage, marges, 2 colonnes
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "2")

    # Titre
    doc.add_heading(f"QCM - Sujet {numero_sujet}", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("Nom / Prénom : ______________________\n")

    # Ajouter chaque question dans un tableau à 1 ligne 1 cellule (anti-coupure)
    for i, q in enumerate(questions):
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = f"{i+1}. {q['texte']}\n"
        for j, rep in enumerate(q["reponses"]):
            cell.text += f"   {chr(97 + j)}) {rep}\n"
        table.allow_autofit = True
        table.style = 'Table Grid'  # ou enlève cette ligne pour invisible
        doc.add_paragraph("")  # espace léger

    # Cartouche réponses à la fin
    doc.add_page_break()
    doc.add_paragraph(f"Sujet {numero_sujet} Réponses")
    doc.add_paragraph("1 2 3 4 5")
    doc.add_paragraph("6 7 8 9 10")

    doc.save(nom_fichier)
    print(f"📄 Sujet {numero_sujet} exporté sans coupures : {nom_fichier}")