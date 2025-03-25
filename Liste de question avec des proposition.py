from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

questions = [
    ("Quelle est la couleur du ciel par beau temps ?", ["Bleu", "Rouge", "Vert", "Jaune"], "1"),
    ("Combien font 2 + 2 ?", ["3", "4", "5", "6"], "2"),
    ("Quel est le plus grand océan du monde ?", ["Atlantique", "Pacifique", "Indien", "Arctique"], "2"),
    ("Quelle planète est surnommée la planète rouge ?", ["Terre", "Mars", "Jupiter", "Vénus"], "2"),
    ("Quel animal aboie ?", ["Chat", "Chien", "Cheval", "Souris"], "2"),
    ("Quel est le nombre de continents sur Terre ?", ["4", "5", "6", "7"], "4"),
    ("Combien de côtés à un triangle ?", ["2", "3", "4", "5"], "2"),
    ("Quelle est la capitale de l'Allemagne ?", ["Madrid", "Berlin", "Rome", "Paris"], "2"),
    ("Quel est l'animal le plus rapide sur terre ?", ["Guépard", "Lion", "Éléphant", "Tortue"], "1"),
    ("Combien de mois dans une année ?", ["10", "11", "12", "13"], "3"),
    ("Quelle est la boisson la plus consommée dans le monde après l'eau ?", ["Coca-Cola", "Thé", "Café", "Jus d'orange"], "2"),
    ("Qui a peint la Joconde ?", ["Vincent van Gogh", "Pablo Picasso", "Léonard de Vinci", "Claude Monet"], "3"),
    ("Quel est le métal principal dans les pièces de 1€ ?", ["Fer", "Or", "Cuivre", "Nickel"], "4"),
    ("Quelle est la capitale du Canada ?", ["Toronto", "Montréal", "Ottawa", "Vancouver"], "3"),
    ("Quel est l'animal emblème de la Chine ?", ["Tigre", "Dragon", "Panda", "Serpent"], "3"),
    ("Combien de doigts à une main humaine ?", ["3", "4", "5", "6"], "3"),
    ("Quel est le gaz le plus abondant dans l'atmosphère terrestre ?", ["Oxygène", "Dioxyde de carbone", "Azote", "Hydrogène"], "3"),
    ("Qui a écrit 'Les Misérables' ?", ["Victor Hugo", "Molière", "Albert Camus", "Jules Verne"], "1"),
    ("Quel est le sport national du Japon ?", ["Football", "Sumo", "Baseball", "Judo"], "2"),
    ("Quelle est la capitale de l'Espagne ?", ["Lisbonne", "Madrid", "Barcelone", "Séville"], "2")
]

# Reformatage
questions_formatees = [{"texte": q[0], "reponses": q[1], "bonne": q[2]} for q in questions]

def generer_sujet_docx(questions, numero_sujet, nom_fichier):
    doc = Document()

    # Mise en page paysage
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "2")

    # Titre
    titre = doc.add_heading(f"QCM - Sujet {numero_sujet}", level=1)
    titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("Nom / Prénom : ____________________________\n")

    # Questions
    for i, q in enumerate(questions):
        p = doc.add_paragraph()
        run_q = p.add_run(f"{i+1}. {q['texte']}")
        run_q.bold = True
        p.space_after = Pt(4)

        for j, rep in enumerate(q["reponses"]):
            p2 = doc.add_paragraph(f"{chr(97 + j)}) {rep}", style='List Bullet')
            p2.paragraph_format.left_indent = Inches(0.3)
            p2.paragraph_format.space_after = Pt(2)

    from docx.shared import Cm

    # Titre de la grille
    doc.add_paragraph("\nGrille de réponses :")


    # Ajouter un paragraphe centré pour le tableau
    para = doc.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Créer un tableau dans ce paragraphe
    table = doc.add_table(rows=2, cols=11)
    table.style = 'Table Grid'
    table_width = section.page_width / 2  # La moitié de la largeur de la page
    for row in table.rows:
        for cell in row.cells:
            cell.width = table_width / len(row.cells)  # Ajuste chaque cellule pour que le tableau soit équilibré

    # Réduire la largeur de chaque cellule
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(1.2)  # Environ la moitié de largeur normale

    # Remplir les cellules
    question_counter = 1
    for row in range(2):
        for col in range(11):
            cell = table.cell(row, col)
            if col == 5:
                cell.text = "-"
            else:
                if question_counter <= 20:
                    cell.text = str(question_counter)
                    question_counter += 1

    # Corrigé
    doc.add_paragraph("\nCorrigé :").runs[0].bold = True
    doc.add_paragraph(' '.join([q["bonne"] for q in questions]))

    doc.save(nom_fichier)
    print(f"📄 Sujet {numero_sujet} exporté avec style propre : {nom_fichier}")

# Génération
generer_sujet_docx(questions_formatees, numero_sujet=1, nom_fichier="QCM_Sujet_1.docx")

